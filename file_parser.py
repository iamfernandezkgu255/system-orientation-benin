import pandas as pd
import docx
import io
from typing import List, Dict, Any, Optional
import streamlit as st

class FileParser:
    """
    Classe pour parser les fichiers Excel (.xlsx) et Word (.docx) contenant les données des étudiants.
    """
    
    def __init__(self):
        self.required_columns = [
            'Nom', 'Prénom', 'Date de Naissance', 
            'Lieu de Naissance', 'Filière Actuelle', 'Carrière Envisagée'
        ]
    
    def parse_file(self, uploaded_file) -> List[Dict[str, Any]]:
        """
        Parse un fichier téléversé et retourne une liste de dictionnaires représentant les étudiants.
        
        Args:
            uploaded_file: Fichier téléversé via Streamlit
            
        Returns:
            List[Dict[str, Any]]: Liste des données étudiants
        """
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        try:
            if file_extension == 'xlsx':
                return self._parse_excel(uploaded_file)
            elif file_extension == 'docx':
                return self._parse_word(uploaded_file)
            else:
                raise ValueError(f"Format de fichier non supporté: {file_extension}")
        except Exception as e:
            st.error(f"Erreur lors du parsing du fichier: {str(e)}")
            return []
    
    def _parse_excel(self, file) -> List[Dict[str, Any]]:
        """
        Parse un fichier Excel.
        
        Args:
            file: Fichier Excel
            
        Returns:
            List[Dict[str, Any]]: Données des étudiants
        """
        try:
            # Lire le fichier Excel
            df = pd.read_excel(file, engine='openpyxl')
            
            # Nettoyer les noms de colonnes (supprimer les espaces en début/fin)
            df.columns = df.columns.str.strip()
            
            # Vérifier la présence des colonnes requises
            missing_columns = self._check_required_columns(df.columns.tolist())
            if missing_columns:
                raise ValueError(f"Colonnes manquantes dans le fichier Excel: {', '.join(missing_columns)}")
            
            # Convertir en liste de dictionnaires
            students_data = []
            for index, row in df.iterrows():
                student_data = {}
                for col in self.required_columns:
                    value = row[col]
                    # Convertir les valeurs NaN en chaînes vides
                    student_data[col] = str(value) if pd.notna(value) else ""
                
                # Valider que l'étudiant a au moins un nom et une filière
                if student_data['Nom'].strip() and student_data['Filière Actuelle'].strip():
                    students_data.append(student_data)
                else:
                    st.warning(f"Ligne {index + 2} ignorée: Nom ou Filière Actuelle manquant")
            
            return students_data
            
        except Exception as e:
            raise Exception(f"Erreur lors de la lecture du fichier Excel: {str(e)}")
    
    def _parse_word(self, file) -> List[Dict[str, Any]]:
        """
        Parse un fichier Word contenant des données structurées.
        Attendu: un tableau avec les colonnes requises ou un format texte structuré.
        
        Args:
            file: Fichier Word
            
        Returns:
            List[Dict[str, Any]]: Données des étudiants
        """
        try:
            # Lire le document Word
            doc = docx.Document(file)
            students_data = []
            
            # Essayer d'abord de lire les tableaux
            if doc.tables:
                students_data = self._parse_word_tables(doc.tables)
            
            # Si aucun tableau ou données vides, essayer de parser le texte
            if not students_data:
                students_data = self._parse_word_text(doc)
            
            return students_data
            
        except Exception as e:
            raise Exception(f"Erreur lors de la lecture du fichier Word: {str(e)}")
    
    def _parse_word_tables(self, tables) -> List[Dict[str, Any]]:
        """
        Parse les tableaux dans un document Word.
        
        Args:
            tables: Liste des tableaux du document
            
        Returns:
            List[Dict[str, Any]]: Données des étudiants
        """
        students_data = []
        
        for table in tables:
            if len(table.rows) < 2:  # Besoin d'au moins une ligne d'en-tête et une ligne de données
                continue
            
            # Extraire les en-têtes (première ligne)
            headers = []
            header_row = table.rows[0]
            for cell in header_row.cells:
                headers.append(cell.text.strip())
            
            # Vérifier si ce tableau contient les colonnes requises
            missing_columns = self._check_required_columns(headers)
            if missing_columns:
                st.warning(f"Tableau ignoré - colonnes manquantes: {', '.join(missing_columns)}")
                continue
            
            # Extraire les données
            for row_idx, row in enumerate(table.rows[1:], start=2):
                student_data = {}
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(headers):
                        header = headers[col_idx]
                        if header in self.required_columns:
                            student_data[header] = cell.text.strip()
                
                # Valider que l'étudiant a au moins un nom et une filière
                if (student_data.get('Nom', '').strip() and 
                    student_data.get('Filière Actuelle', '').strip()):
                    # S'assurer que toutes les colonnes requises sont présentes
                    for col in self.required_columns:
                        if col not in student_data:
                            student_data[col] = ""
                    students_data.append(student_data)
                else:
                    st.warning(f"Ligne {row_idx} du tableau ignorée: Nom ou Filière Actuelle manquant")
        
        return students_data
    
    def _parse_word_text(self, doc) -> List[Dict[str, Any]]:
        """
        Parse le texte libre d'un document Word.
        Recherche des patterns comme "Nom: John Doe" etc.
        
        Args:
            doc: Document Word
            
        Returns:
            List[Dict[str, Any]]: Données des étudiants
        """
        students_data = []
        current_student = {}
        
        # Patterns de recherche pour les champs
        field_patterns = {
            'Nom': ['nom:', 'nom :', 'name:'],
            'Prénom': ['prénom:', 'prénom :', 'prenom:', 'firstname:'],
            'Date de Naissance': ['date de naissance:', 'naissance:', 'né le:', 'birthdate:'],
            'Lieu de Naissance': ['lieu de naissance:', 'lieu:', 'né à:', 'birthplace:'],
            'Filière Actuelle': ['filière:', 'filiere:', 'filière actuelle:', 'field:'],
            'Carrière Envisagée': ['carrière:', 'carriere:', 'carrière envisagée:', 'career:']
        }
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip().lower()
            
            if not text:
                # Paragraphe vide - potentiellement fin d'un étudiant
                if self._is_student_complete(current_student):
                    students_data.append(current_student.copy())
                    current_student = {}
                continue
            
            # Rechercher les champs dans le texte
            for field, patterns in field_patterns.items():
                for pattern in patterns:
                    if pattern in text:
                        # Extraire la valeur après le pattern
                        start_idx = text.find(pattern) + len(pattern)
                        value = paragraph.text[start_idx:].strip()
                        current_student[field] = value
                        break
        
        # Ajouter le dernier étudiant s'il est complet
        if self._is_student_complete(current_student):
            students_data.append(current_student)
        
        return students_data
    
    def _check_required_columns(self, columns: List[str]) -> List[str]:
        """
        Vérifie la présence des colonnes requises.
        
        Args:
            columns: Liste des colonnes présentes
            
        Returns:
            List[str]: Liste des colonnes manquantes
        """
        missing_columns = []
        for required_col in self.required_columns:
            if required_col not in columns:
                # Essayer de trouver des variations (insensible à la casse, avec/sans accents)
                found = False
                for col in columns:
                    if self._normalize_column_name(col) == self._normalize_column_name(required_col):
                        found = True
                        break
                if not found:
                    missing_columns.append(required_col)
        
        return missing_columns
    
    def _normalize_column_name(self, name: str) -> str:
        """
        Normalise un nom de colonne pour la comparaison.
        
        Args:
            name: Nom de la colonne
            
        Returns:
            str: Nom normalisé
        """
        import unicodedata
        # Supprimer les accents et convertir en minuscules
        normalized = unicodedata.normalize('NFD', name.lower())
        normalized = ''.join(c for c in normalized if unicodedata.category(c) != 'Mn')
        # Supprimer les espaces et caractères spéciaux
        normalized = ''.join(c for c in normalized if c.isalnum())
        return normalized
    
    def _is_student_complete(self, student_data: Dict[str, str]) -> bool:
        """
        Vérifie si un étudiant a les données minimales requises.
        
        Args:
            student_data: Données de l'étudiant
            
        Returns:
            bool: True si les données sont suffisantes
        """
        return (student_data.get('Nom', '').strip() and 
                student_data.get('Filière Actuelle', '').strip())
    
    def validate_student_data(self, students_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Valide les données des étudiants et retourne un rapport de validation.
        
        Args:
            students_data: Liste des données étudiants
            
        Returns:
            Dict[str, Any]: Rapport de validation
        """
        validation_report = {
            'total_students': len(students_data),
            'valid_students': 0,
            'warnings': [],
            'errors': []
        }
        
        for i, student in enumerate(students_data):
            student_id = f"Étudiant {i+1} ({student.get('Nom', 'N/A')} {student.get('Prénom', 'N/A')})"
            
            # Vérifications de base
            if not student.get('Nom', '').strip():
                validation_report['errors'].append(f"{student_id}: Nom manquant")
                continue
            
            if not student.get('Filière Actuelle', '').strip():
                validation_report['errors'].append(f"{student_id}: Filière Actuelle manquante")
                continue
            
            if not student.get('Carrière Envisagée', '').strip():
                validation_report['warnings'].append(f"{student_id}: Carrière Envisagée manquante")
            
            validation_report['valid_students'] += 1
        
        return validation_report